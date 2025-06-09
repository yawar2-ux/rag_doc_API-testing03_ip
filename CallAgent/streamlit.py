import os
from openai import AzureOpenAI
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from textblob import TextBlob
import json
import speech_recognition as sr
from fastapi import UploadFile
import docx
from CallAgent.db import get_db_connection, generate_cust_id
from CallAgent.db import CallCenterDB, validate_mobile_no
import pdfplumber
import chromadb
from chromadb.config import Settings
import uuid
import scipy
import pandas as pd
from typing import Optional, Dict, List
import logging
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import base64
import io
from collections import Counter
import re

class WordCloudGenerator:
    def __init__(self):
        self.customer_text = []  # Only store customer text
        self.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.customer_word_frequencies = Counter()

    def reset(self):
        """Reset for new call session"""
        self.customer_text = []
        self.customer_word_frequencies.clear()
        self.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")

    def update_customer_text(self, customer_text):
        """Add only customer text"""
        self.customer_text.append(customer_text)
        words = re.findall(r'\b\w+\b', customer_text.lower())
        self.customer_word_frequencies.update(words)

    def generate_wordcloud(self):
        """Generate wordcloud from customer text only"""
        if not self.customer_text:
            return None

        current_text = " ".join(self.customer_text)
        if not current_text.strip():
            return None

        wordcloud = WordCloud(
            width=800,
            height=400,
            background_color='white',
            max_words=100
        ).generate_from_frequencies(self.customer_word_frequencies)

        filename = f"customer_wordcloud_{self.session_id}.png"
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis('off')
        plt.savefig(filename, bbox_inches='tight', pad_inches=0)
        plt.close()

        with open(filename, 'rb') as img_file:
            base64_image = base64.b64encode(img_file.read()).decode()

        return base64_image


class CallCenterAssistant:
    def __init__(self):
        super().__init__()
        self.azure_client = AzureOpenAI(
            azure_endpoint="https://rahul.openai.azure.com/",
            api_key="8b1d436e85d1452bbcbfd5905921efa6",
            api_version="2024-02-15-preview"
        )
        self.deployment_name = "RAG"
        self.call_transcript = []
        self.document_store = {}
        self.vectorizer = TfidfVectorizer(
            stop_words='english',
            max_features=10000,
            ngram_range=(1, 2)
        )
        self.vectors = None
        self.contents = []
        self.db = get_db_connection()
        self.db = CallCenterDB()
        self.start_time = None
        self.current_cust_id = None
        self.current_agent_id = None
        self.sentiment_scores = []
        self.wordcloud_generator = WordCloudGenerator()
        self.current_mobile_no = None
        self.previous_summary = None
        self.vectorizer = TfidfVectorizer()
        self.chroma_client = chromadb.PersistentClient(path="./chroma_storage")
        self.qa_collection = self.chroma_client.get_or_create_collection(
            name="qa_history",
            metadata={"hnsw:space": "cosine"}
        )
        self.document_chunks = {}  # Store text chunks with doc_ids
        self.chunk_size = 1000
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)


    def search_chroma_history(self, query, k=1):
        """Search ChromaDB for similar queries"""
        results = self.collection.query(
            query_texts=[query],
            n_results=k
        )
        if results and results['distances'][0][0] < 0.3:  # Similarity threshold
            return results['metadatas'][0][0]
        return None

    def store_in_chroma(self, query, response, source):
        """Store Q&A pair in ChromaDB"""
        self.collection.add(
            documents=[query],
            metadatas=[{
                "response": response,
                "source": source,
                "timestamp": datetime.now().isoformat()
            }],
            ids=[str(uuid.uuid4())]
        )

    def unified_search(self, query):
        """Search across all knowledge sources"""
        # 1. Check document store
        doc_content = self.find_similar_content(query)
        if doc_content:
            return {"content": doc_content, "source": "documents"}

        # 2. Check ChromaDB history
        history_match = self.search_chroma_history(query)
        if history_match:
            return {"content": history_match['response'], "source": "history"}

        # 3. Generate new response with LLM
        llm_response = self.generate_ai_response(query)
        self.store_in_chroma(query, llm_response, "llm")
        return {"content": llm_response, "source": "llm"}


    def get_mobile_number(self):
        while True:
            try:
                mobile_no = input("Please enter 4-digit mobile number: ")
                return validate_mobile_no(mobile_no)
            except ValueError as e:
                print(f"Error: {e}. Please try again.")


    def check_previous_calls(self, current_text, mobile_no):
        """Check for similar previous calls"""
        try:
            # Get previous call summary from DB
            previous_summary = self.db.get_call_history(mobile_no)

            if not previous_summary:
                return None

            # Calculate similarity
            texts = [previous_summary, current_text]
            tfidf_matrix = self.vectorizer.fit_transform(texts)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

            # Return summary if similarity threshold met
            if similarity > 0.3:  # Adjust threshold as needed
                return previous_summary
            return None

        except Exception as e:
            print(f"Error checking previous calls: {e}")
            return None


    async def process_multiple_documents(self, files: List[UploadFile]):
        """Process multiple banking documents concurrently."""
        try:
            if not files:
                return {
                    'status': 'error',
                    'message': 'No files provided'
                }

            # Track results for each file
            results = {
                'successful': [],
                'failed': [],
                'total': len(files)
            }

            # Process files concurrently
            import asyncio
            tasks = [self.process_document(file) for file in files]
            processed = await asyncio.gather(*tasks, return_exceptions=True)

            # Compile results
            for file, result in zip(files, processed):
                if isinstance(result, Exception):
                    results['failed'].append({
                        'filename': file.filename,
                        'error': str(result)
                    })
                elif result['status'] == 'success':
                    results['successful'].append({
                        'filename': file.filename,
                        'doc_id': result['doc_id'],
                        'chunks': result['chunks'],
                        'tables': result['tables']
                    })
                else:
                    results['failed'].append({
                        'filename': file.filename,
                        'error': result['message']
                    })

            # Calculate stats
            results['stats'] = {
                'total_processed': len(results['successful']),
                'total_failed': len(results['failed']),
                'success_rate': f"{(len(results['successful'])/len(files))*100:.1f}%"
            }

            return {
                'status': 'success' if results['successful'] else 'partial_success' if results['failed'] else 'error',
                'results': results
            }

        except Exception as e:
            self.logger.error(f"Multiple document processing error: {str(e)}")
            return {
                'status': 'error',
                'message': f'Failed to process documents: {str(e)}'
            }


    def _chunk_text(self, text: str, chunk_size: int = 1000) -> list:
        """Split text into overlapping chunks for better context"""
        try:
            if not text:
                return []

            words = text.split()
            chunks = []

            # Create overlapping chunks
            for i in range(0, len(words), chunk_size // 2):
                chunk = ' '.join(words[i:i + chunk_size])
                if chunk:
                    chunks.append(chunk)

            return chunks

        except Exception as e:
            self.logger.error(f"Error chunking text: {str(e)}")
            return []


    def validate_document(self, content: str) -> tuple[bool, str]:
        """Validate document is English and banking-related."""
        try:
            # Check if content is empty
            if not content.strip():
                return False, "Document is empty."

            # Check if content is primarily English (basic ASCII check)
            def is_english_content(content):
                ascii_chars = sum(c.isascii() for c in content)
                return ascii_chars / len(content) >= 0.85

            if not is_english_content(content):
                return False, "Document must primarily contain English text."

            # Check for banking keywords
            def contains_banking_keywords(content):
                keywords = [
                    'account', 'bank', 'credit', 'debit', 'loan', 'deposit',
                    'payment', 'transaction', 'interest', 'savings', 'finance',
                    'transfer', 'balance', 'card', 'statement', 'atm'
                ]
                text_lower = content.lower()
                return sum(1 for word in keywords if word in text_lower) >= 2

            if not contains_banking_keywords(content):
                return False, "Document must include sufficient banking-related keywords."

            return True, "Valid document."
        except Exception as e:
            return False, f"Validation error: {str(e)}"


    async def process_document(self, files: List[UploadFile]):
        """Process multiple documents and store their content."""
        try:
            for file in files:
                file_ext = file.filename.rsplit('.', 1)[1].lower()
                file_content = await file.read()
                content = ""
                tables = []

                if file_ext == 'pdf':
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        for page in pdf.pages:
                            text = page.extract_text()
                            if text:
                                content += text
                            page_tables = page.extract_tables()
                            if page_tables:
                                for table in page_tables:
                                    if table and any(table):
                                        try:
                                            df = pd.DataFrame(table)
                                            tables.append(df)
                                        except Exception as table_error:
                                            self.logger.error(f"Table processing error: {str(table_error)}")
                                            continue

                elif file_ext == 'docx':
                    doc = docx.Document(io.BytesIO(file_content))
                    content = '\n'.join(p.text for p in doc.paragraphs if p.text)
                    for table in doc.tables:
                        if table.rows:
                            try:
                                data = [[cell.text for cell in row.cells] for row in table.rows]
                                if data:
                                    df = pd.DataFrame(data[1:], columns=data[0])
                                    tables.append(df)
                            except Exception as table_error:
                                self.logger.error(f"Table processing error: {str(table_error)}")

                elif file_ext == 'txt':
                    content = file_content.decode('utf-8', errors='ignore')

                elif file_ext == 'csv':
                    df = pd.read_csv(io.StringIO(file_content.decode('utf-8', errors='ignore')))
                    content = df.to_string()
                    tables.append(df)

                else:
                    raise ValueError(f"Unsupported file type: {file_ext}")

                doc_id = str(len(self.document_store) + 1)
                self.document_store[doc_id] = {
                    'content': content,
                    'tables': tables
                }

            # Update vectors for document search
            self.vectors = self.vectorizer.fit_transform([doc['content'] for doc in self.document_store.values()])
            return {'status': 'success', 'doc_ids': list(self.document_store.keys())}

        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    def search_documents(self, query: str) -> Optional[Dict[str, str]]:
        """Search for a query in the processed documents."""
        if not self.document_store:
            return None

        query_vector = self.vectorizer.transform([query])
        similarities = cosine_similarity(query_vector, self.vectors)[0]
        best_idx = similarities.argmax()

        if similarities[best_idx] > 0.085:  # Adjust threshold for relevance
            doc_id = list(self.document_store.keys())[best_idx]
            best_match = self.document_store[doc_id]['content']

            # Search in tables for more specific information
            table_matches = [] 
            for table in self.document_store[doc_id]['tables']:
                if query.lower() in table.to_string().lower():
                    table_matches.append(table.to_string())

            return {
                "content": best_match,
                "tables": "\n\n".join(table_matches) if table_matches else None
            }
        return None


    def transcribe_audio(self):
        recognizer = sr.Recognizer()
        with sr.Microphone() as source:
            print("Listening...")
            audio = recognizer.listen(source, timeout=5)
            try:
                text = recognizer.recognize_google(audio)
                print(f"Customer: {text}")

                # Add to transcript and calculate sentiment
                self.call_transcript.append(f"Customer: {text}")
                sentiment = self.analyze_sentiment(text)
                print(f"Sentiment Analysis: {sentiment}")

                # Parse sentiment score and add to sentiment_scores
                sentiment_data = json.loads(sentiment)
                score = float(sentiment_data.get('sentiment_score', 50))  # Default to 50 if parsing fails
                self.sentiment_scores.append(score)

                return text
            except sr.UnknownValueError:
                print("Could not understand audio.")
                return ""
            except sr.RequestError as e:
                print(f"Speech recognition service error: {e}")
                return ""

    def compute_average_sentiment(self):
        """Compute the average sentiment from all recorded scores."""
        if not self.sentiment_scores:
            return 50.0  # Default neutral sentiment
        return sum(self.sentiment_scores) / len(self.sentiment_scores)  # Returns a single float


    def end_call_and_save(self):
        """End the call, compute average sentiment, and save the call record."""
        try:
            end_time = datetime.now()
            avg_sentiment = self.compute_average_sentiment()  # Compute the average sentiment
            print(f"Final Average Sentiment: {avg_sentiment}, Type: {type(avg_sentiment)}")  # Debug print

            summary = self.generate_summary()
            print(f"Call Summary: {summary}")

            # Save the call record
            success = self.db.save_call_record(
                self.start_time,
                self.current_cust_id,
                self.current_agent_id,
                self.current_mobile_no,
                self.call_transcript,
                summary,
                avg_sentiment  # Ensure this is a single float value
            )
            if success:
                print("Call record saved successfully!")
            else:
                print("Failed to save call record.")
        except Exception as e:
            print(f"Error ending call: {e}")


    def analyze_sentiment(self, text):
            try:
                # Use TextBlob for simple sentiment analysis
                analysis = TextBlob(text)
                sentiment_score = (analysis.sentiment.polarity + 1) * 50  # Convert polarity (-1 to 1) to score (0 to 100)

                if sentiment_score > 60:
                    sentiment_category = "Positive"
                elif sentiment_score < 40:
                    sentiment_category = "Negative"
                else:
                    sentiment_category = "Neutral"

                response = json.dumps({
                    "sentiment_score": sentiment_score,
                    "sentiment_category": sentiment_category,
                    "explanation": f"The sentiment is categorized as {sentiment_category} based on the sentiment score."
                })
                return response
            except Exception as e:
                self.logger.error(f"Error analyzing sentiment: {e}")
                return json.dumps({
                    "sentiment_score": 50,  # Default neutral
                    "sentiment_category": "Neutral",
                    "explanation": "Error in sentiment analysis"
                })
    def analyze_sentiment(self, text):
        """
        Analyzes customer sentiment using OpenAI Azure LLM and returns a JSON-formatted string.
        The response includes sentiment_score, sentiment_category, and an explanation.
        """
        try:
            # Call Azure LLM for sentiment analysis
            response = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an intelligent AI for sentiment analysis. Analyze the sentiment of the following text and "
                            "provide a JSON response with:\n"
                            "1. 'sentiment_score' (0-100 scale), where 0 is very negative and 100 is very positive.\n"
                            "2. 'sentiment_category' (Positive, Neutral, or Negative).\n"
                            "3. 'explanation' for why the sentiment was categorized this way."
                        )
                    },
                    {"role": "user", "content": text}
                ],
                temperature=0.3,
                max_tokens=200
            )

            # Parse the response content
            raw_sentiment_data = response.choices[0].message.content

            # Check if the response is a string (likely JSON text) and parse it
            if isinstance(raw_sentiment_data, str):
                sentiment_data = json.loads(raw_sentiment_data)
            else:
                raise ValueError("Unexpected response format from Azure LLM.")

            # Extract fields from the response
            sentiment_score = sentiment_data.get("sentiment_score", 50)  # Default to 50 if missing
            sentiment_category = sentiment_data.get("sentiment_category", "Neutral")
            explanation = sentiment_data.get("explanation", "No explanation provided.")

            # Construct the response in the desired format
            return json.dumps({
                "sentiment_score": sentiment_score,
                "sentiment_category": sentiment_category,
                "explanation": explanation
            })
        except Exception as e:
            self.logger.error(f"Error analyzing sentiment with LLM: {e}")
            return json.dumps({
                "sentiment_score": 50,  # Default neutral
                "sentiment_category": "Neutral",
                "explanation": "Error in sentiment analysis with LLM"
            })            


    def generate_ai_response(self, text):
        """Generates AI response to comprehensive banking and customer service queries."""
        banking_keywords = [
            'account', 'bank', 'credit', 'debit', 'loan', 'deposit', 'transfer',
            'payment', 'balance', 'card', 'statement', 'interest', 'atm',
            'transaction', 'banking', 'check', 'savings', 'withdraw',
            'insurance', 'policy', 'premium', 'claim', 'coverage', 'insure',
            'trading', 'stocks', 'shares', 'mutual fund', 'portfolio', 'investment',
            'dividend', 'equity', 'securities', 'market',
            'branch', 'locker', 'forex', 'exchange', 'currency', 'remittance',
            'draft', 'challan', 'passbook', 'nominee', 'kyc', 'income', 'tax', 'economy'
        ]

        completion = self.azure_client.chat.completions.create(
            model=self.deployment_name,
            messages=[
                {"role": "system", "content": (
                    "You are a professional AI banking assistant and call center agent. You handle a wide variety of customer service requests related to banking, including:\n\n"
                    "1. Account Services:\n"
                    "   - Account opening, closing, and statements\n"
                    "   - Savings, checking, and fixed deposits\n"
                    "   - Account balance and transactions\n\n"
                    "2. Banking Products:\n"
                    "   - Credit cards and debit cards\n"
                    "   - Issues related to credit or debit cards\n"
                    "   - Loans (personal, home, auto, education)\n"
                    "   - Investment products\n\n"
                    "3. Digital Banking:\n"
                    "   - Online and mobile banking\n"
                    "   - Digital payments and transfers\n"
                    "   - Electronic statements\n\n"
                    "4. Insurance Services:\n"
                    "   - Life insurance\n"
                    "   - Health insurance\n"
                    "   - Property insurance\n"
                    "   - Vehicle insurance\n\n"
                    "5. Trading & Investments:\n"
                    "   - Stock trading\n"
                    "   - Mutual funds\n"
                    "   - Portfolio management\n"
                    "   - Investment advisory\n\n"
                    "6. Retail Banking:\n"
                    "   - Branch services\n"
                    "   - ATM services\n"
                    "   - Safe deposit lockers\n"
                    "   - Foreign exchange\n\n"
                    "7. Customer Support:\n"
                    "   - PIN/password reset\n"
                    "   - Card blocking or replacement\n"
                    "   - Banking charges and fees\n\n"
                    "8. Security:\n"
                    "   - Fraud prevention\n"
                    "   - Account security\n"
                    "   - Safe banking practices\n\n"
                    "Important Rules:\n"
                    "- Only respond to customer-related queries but based on banks.\n"
                    "- For non-banking questions, respond: 'I can only assist with Customer-related inquiries. Please ask a relevant question.'\n"
                    "- For customer service-related queries, such as connecting to specific departments, acknowledge the request and offer assistance, e.g., 'Sure, I will do that for you. Is there anything else I can assist you with?'\n"
                    "- Maintain a helpful, empathetic, and professional tone at all times.\n"
                    "- For ambiguous questions (e.g., 'Can you help me?'), ask clarifying questions to better understand the customer’s needs.\n"
                    "- Reject inappropriate or abusive content politely.\n"
                    "- Never assist with fraud, money laundering, or unauthorized access.\n"
                    "- Protect customer privacy by avoiding requests for sensitive information.\n"
                    "- Always confirm if the customer needs further assistance before ending the conversation.\n\n"
                    "When answering, ensure your response is concise, directly addresses the query, and provides appropriate next steps or resolutions as needed.\n\n"
                    "Keywords you can consider but optinal: " + ', '.join(banking_keywords)
                )},
                {"role": "user", "content": text}
            ],
            temperature=0.7,
            max_tokens=300
        )
        return completion.choices[0].message.content


    def generate_response(self, query: str) -> str:
        """Generate response for banking queries only"""
        if not self.validate_banking_query(query):
            return "I can only assist with banking-related inquiries. Please ask a banking question."

        self.logger.info(f"Processing banking query: {query}")

        # First check banking documents
        doc_content = self.find_similar_content(query)

        if doc_content:
            self.logger.info("Found relevant banking document content")
            return self.generate_ai_response(
                f"Based on our banking documentation: {doc_content}\n\nQuery: {query}"
            )

        self.logger.info("No relevant document content found, generating AI response")
        return self.generate_ai_response(query)

    def generate_summary(self):
        """Generates a summary of the entire call."""
        conversation = "\n".join(self.call_transcript)
        completion = self.azure_client.chat.completions.create(
            model=self.deployment_name,
            messages=[
                {"role": "system", "content": (
                    "You are an AI assistant summarizing customer service calls. "
                    "Summarize the following conversation into key points, highlighting important details, concerns, and resolutions."
                )},
                {"role": "user", "content": conversation}
            ],
            temperature=0.5,
            max_tokens=500
        )
        return completion.choices[0].message.content

    def analyze_sales_opportunity(self, conversation_history):
        """Analyzes the conversation for potential sales opportunities."""
        try:
            # More robust conversation text building
            conversation_text = ""
            for entry in conversation_history:
                speaker = entry.get('type', 'unknown').title()
                message = entry.get('message', '')
                conversation_text += f"{speaker}: {message}\n"

            completion = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": """
                        You are a sales opportunity analyzer. Based on the conversation, identify:
                        1) Current products/services mentioned
                        2) Customer pain points and needs
                        3) Potential upsell/cross-sell opportunities
                        4) Customer's buying signals or interest level
                        5) Probability of conversion (High/Medium/Low)

                        Return the analysis in this JSON format:
                        {
                            "opportunities": ["list of specific products/services to offer"],
                            "pain_points": ["list of customer needs/issues"],
                            "buying_signals": ["list of positive buying indicators"],
                            "next_steps": ["recommended sales actions"],
                            "probability": "High/Medium/Low",
                            "explanation": "Brief explanation of the analysis"
                        }
                    """},
                    {"role": "user", "content": conversation_text}
                ],
                temperature=0.7,
                max_tokens=800,
                response_format={"type": "json_object"}
            )

            return json.loads(completion.choices[0].message.content)

        except Exception as e:
            print(f"Error in sales opportunity analysis: {str(e)}")
            return {
                "opportunities": [],
                "pain_points": [],
                "buying_signals": [],
                "next_steps": [],
                "probability": "Low",
                "explanation": f"Error analyzing sales opportunities: {str(e)}"
            }

    def run(self):
        """Main call center session handler"""
        print("Call Center Assistant is running...")
        self.wordcloud_generator = WordCloudGenerator()
        self.wordcloud_generator.reset()
        session_id = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Initialize previous summary from DB
        self.previous_summary = self.db.get_call_history(self.current_mobile_no)
        if self.previous_summary:
            print("\nPrevious call summary found:")
            print(self.previous_summary)

        while True:
            customer_query = self.transcribe_audio()
            if not customer_query:
                continue

            # Check similarity with previous calls
            if self.previous_summary:
                texts = [self.previous_summary, customer_query]
                tfidf_matrix = self.vectorizer.fit_transform(texts)
                similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

                if similarity > 0.3:
                    print("\nSimilar issue identified from previous call:")
                    print(self.previous_summary)

            # Update transcript and customer-only wordcloud
            self.call_transcript.append(f"Customer: {customer_query}")
            self.wordcloud_generator.update_customer_text(customer_query)

            # Generate AI response
            ai_response = self.generate_ai_response(customer_query)
            print(f"Assistant: {ai_response}")

            # Update transcript only (not wordcloud)
            self.call_transcript.append(f"Assistant: {ai_response}")

            if "goodbye" in customer_query.lower() or "end call" in customer_query.lower():
                # Generate final customer-only wordcloud
                final_wordcloud = self.wordcloud_generator.generate_wordcloud()
                summary = self.generate_summary()

                # Save customer wordcloud
                try:
                    wordcloud_file = f"customer_wordcloud_{session_id}.png"
                    with open(wordcloud_file, 'wb') as f:
                        f.write(base64.b64decode(final_wordcloud))
                    print(f"\nCustomer wordcloud saved as: {wordcloud_file}")
                except Exception as e:
                    print(f"Error saving wordcloud: {e}")

                print("\nCall Summary:")
                print(summary)

                return {
                    'summary': summary,
                    'wordcloud': final_wordcloud,
                    'session_id': session_id,
                    'mobile_no': self.current_mobile_no,
                    'transcript': self.call_transcript
                }



# from call_center_assistant import CallCenterAssistant
import glob
import os

class FileHandler:
    def __init__(self, path):
        self.file = open(path, 'rb')
        self.filename = os.path.basename(path)

    def read(self):
        return self.file.read()

    def close(self):
        self.file.close()

import glob
import os
import logging
from datetime import datetime

def run_call_center_demo():
    # Initialize database and assistant
    db = get_db_connection()
    if not db:
        print("Error: Could not connect to database. Exiting...")
        return

    assistant = CallCenterAssistant()
    assistant.db = db
    assistant.start_time = None
    assistant.current_cust_id = None
    assistant.current_agent_id = None
    assistant.sentiment_scores = []

    while True:
        print("\nDocument Upload Options:")
        print("1. Upload a single document")
        print("2. Upload multiple documents from a directory")
        print("3. Start call center session")
        print("4. Exit")

        choice = input("Enter your choice (1-4): ")

        if choice == "1":
            path = input("Enter document path: ")
            try:
                file_handler = FileHandler(path)
                result = assistant.process_document(file_handler)
                file_handler.close()
                print(f"Document processed: {result['status']}")
            except Exception as e:
                print(f"Error processing document: {e}")

        elif choice == "2":
            directory = input("Enter directory path containing documents: ")
            file_types = "*.pdf *.txt *.docx *.csv *.html"
            files = []
            for ext in file_types.split():
                files.extend(glob.glob(os.path.join(directory, ext)))

            print(f"\nFound {len(files)} documents:")
            for f in files:
                print(f"- {os.path.basename(f)}")

            confirm = input("\nProcess these documents? (y/n): ")
            if confirm.lower() == 'y':
                for file_path in files:
                    try:
                        file_handler = FileHandler(file_path)
                        result = assistant.process_document(file_handler)
                        file_handler.close()
                        print(f"Processed {os.path.basename(file_path)}: {result['status']}")
                    except Exception as e:
                        print(f"Error processing {file_path}: {e}")

        elif choice == "3":
            try:
                # Get and validate mobile number
                mobile_no = input("Please enter 4-digit mobile number: ")
                mobile_no = validate_mobile_no(mobile_no)
                print("Mobile number validated successfully.")

                # Initialize call session
                assistant.start_time = datetime.now()
                assistant.current_cust_id = generate_cust_id()
                assistant.current_agent_id = db.select_available_agent()
                assistant.current_mobile_no = mobile_no

                print("\nStarting call with:")
                print(f"Mobile: {mobile_no}")
                print(f"Customer ID: {assistant.current_cust_id}")
                print(f"Agent ID: {assistant.current_agent_id}")

                # Start conversation loop
                while True:
                    customer_query = assistant.transcribe_audio()
                    if not customer_query:
                        continue

                    # Search for the most relevant previous summary dynamically
                    similar_summary = db.find_similar_in_customer_history(
                        customer_query,
                        mobile_no
                    )

                    if similar_summary:
                        print("\nRelevant summary from your previous calls:")
                        print(f"{similar_summary}\n")
                    else:
                        print("\nNo relevant summary found in your call history.\n")

                    # Process current conversation
                    assistant.call_transcript.append(f"Customer: {customer_query}")
                    assistant.wordcloud_generator.update_customer_text(customer_query)

                    ai_response = assistant.generate_ai_response(customer_query)
                    print(f"Assistant: {ai_response}")
                    assistant.call_transcript.append(f"Assistant: {ai_response}")

                    # End call condition
                    if "goodbye" in customer_query.lower() or "end call" in customer_query.lower():
                        # End call and save records
                        end_time = datetime.now()
                        summary = assistant.generate_summary()
                        assistant.db.save_call_record(
                            assistant.start_time,
                            assistant.current_cust_id,
                            assistant.current_agent_id,
                            mobile_no,
                            assistant.call_transcript,
                            summary,
                            assistant.sentiment_scores
                        )
                        break

            except ValueError as e:
                print(f"Error: {e}")
            except Exception as e:
                print(f"Error in call session: {e}")

        elif choice == "4":
            print("Exiting...")
            try:
                db.close()
            except:
                pass
            break

        else:
            print("Invalid choice. Please try again.")


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        filename='call_center.log'
    )
    run_call_center_demo()
